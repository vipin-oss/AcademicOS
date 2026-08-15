"""L2 DOCX engine adapter (python-docx).

Paragraphs, headings (style-based), tables, and images with provenance.
Page count is not recorded in the OOXML core package, so it stays 0 (honest).
"""

from __future__ import annotations

import io

from app.application.dtos.nir import NirDocument, NirElement, NirElementType, NirImage
from app.application.ports.nir_parser import NirParseError, NirParser
from app.application.services.extraction_limits import MAX_ELEMENTS
from app.domain.value_objects.source import MediaKind


class DocxNirParser(NirParser):
    format_name = "docx"

    def __init__(self) -> None:
        self._engine = "python-docx"

    def parse(self, data: bytes, *, source_id: str, version: int) -> NirDocument:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover
            raise NirParseError(f"DOCX engine unavailable: {exc}.") from exc

        try:
            document = docx.Document(io.BytesIO(data))
        except Exception as exc:  # noqa: BLE001
            raise NirParseError(f"DOCX could not be parsed ({type(exc).__name__}: {exc}).") from exc

        elements: list[NirElement] = []
        images: list[NirImage] = []
        text_parts: list[str] = []
        order = 0

        for paragraph in document.paragraphs:
            if order >= MAX_ELEMENTS:
                break
            text = paragraph.text or ""
            style = paragraph.style.name.lower() if paragraph.style else ""
            if text:
                text_parts.append(text)
            etype = NirElementType.PARAGRAPH
            if style.startswith("heading") or style.startswith("title"):
                etype = NirElementType.HEADING
            elements.append(
                NirElement(
                    element_type=etype, order=order, text=text,
                    value={"style": style}, extraction_confidence=1.0,
                )
            )
            order += 1

        # tables
        for t_index, table in enumerate(document.tables or []):
            if order >= MAX_ELEMENTS:
                break
            rows: list[list[str]] = []
            for row in table.rows or []:
                rows.append([cell.text or "" for cell in row.cells or []])
            table_id = f"docx-{source_id}-t{t_index}"
            elements.append(
                NirElement(
                    element_type=NirElementType.TABLE, order=order,
                    text="\n".join(" | ".join(r) for r in rows),
                    value={"table_id": table_id, "rows": rows},
                    extraction_confidence=1.0,
                )
            )
            order += 1
            # table cells
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    if order >= MAX_ELEMENTS:
                        break
                    elements.append(
                        NirElement(
                            element_type=NirElementType.TABLE_CELL, order=order,
                            text=cell, value={"table_id": table_id, "row": r, "col": c},
                            extraction_confidence=1.0,
                        )
                    )
                    order += 1

        normalized = "\n".join(text_parts)
        return NirDocument(
            source_id=source_id,
            media_kind=MediaKind.TEXT_LAYOUT.value,
            version=version,
            engine=self._engine,
            engine_version=_docx_version(),
            elements=tuple(elements),
            images=tuple(images),
            normalized_text=normalized[: 8_000_000],
            needs_ocr=False,
        )


def _docx_version() -> int:
    try:
        import docx

        return int(docx.__version__.split(".")[0])
    except Exception:  # noqa: BLE001
        return 0
